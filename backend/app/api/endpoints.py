from fastapi import APIRouter, HTTPException, Request, UploadFile, File
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from pydantic import BaseModel
from typing import Optional, List, Dict
import os
import tempfile
from pathlib import Path
from datetime import datetime

from jinja2 import Template, FileSystemLoader, Environment

from app.agents.text_agent import extract_pdd_sections
from app.agents.diagram_agent import generate_mermaid_diagram
from app.agents.video_agent import (
    transcribe_audio_from_video,
    analyze_video_frames,
    synthesize_video_analysis
)
from app.utils.file_parser import parse_document
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

router = APIRouter()


class GeneratePDDRequest(BaseModel):
    process_text: str


class PDDSection(BaseModel):
    name: str
    content: str


class RefineSectionRequest(BaseModel):
    section_name: str
    current_content: str
    user_feedback: str


class ChatRequest(BaseModel):
    message: str
    context: Optional[str] = None


class ExportPDDRequest(BaseModel):
    process_name: str
    sections: List[Dict[str, str]]
    diagram_code: Optional[str]
    format: str  # 'pdf', 'docx', or 'html'


@router.post("/generate-pdd", response_class=HTMLResponse)
async def generate_pdd(request: Request, body: GeneratePDDRequest):
    """
    Generate a PDD from the provided process text.

    Args:
        request: FastAPI request object
        body: Request body containing process_text

    Returns:
        HTMLResponse with the rendered PDD

    Raises:
        HTTPException: If PDD generation fails
    """
    try:
        # Extract PDD sections using the text agent
        sections = extract_pdd_sections(body.process_text)

        # Get the process name for the title (first section) - strip HTML tags
        import re
        process_name = sections[0]["content"] if sections else "Process Design Document"
        process_name_clean = re.sub('<[^<]+?>', '', process_name).strip()
        process_name = process_name_clean if process_name_clean else "Process Design Document"

        # Generate diagram from process steps (6th section - Detailed Process Steps)
        diagram_code = None
        if len(sections) >= 6:
            process_steps = sections[5]["content"]  # Detailed Process Steps section
            try:
                diagram_code = generate_mermaid_diagram(process_steps)
            except Exception as e:
                print(f"Warning: Failed to generate diagram: {str(e)}")
                diagram_code = None

        # Setup Jinja2 environment
        template_path = Path(__file__).parent.parent / "templates"
        env = Environment(loader=FileSystemLoader(template_path))
        template = env.get_template("pdd_template.html")

        # Render the template
        html_content = template.render(
            process_name=process_name,
            sections=sections,
            diagram_code=diagram_code
        )

        return HTMLResponse(content=html_content)

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate PDD: {str(e)}"
        )


@router.get("/health")
async def health_check():
    """Health check endpoint."""
    return {"status": "healthy"}


@router.post("/upload-and-process", response_class=HTMLResponse)
async def upload_and_process(request: Request, file: UploadFile = File(...)):
    """
    Upload a file (PDF, DOCX, or MP4) and generate a PDD from its content.

    For documents (PDF, DOCX): Extracts text and processes it.
    For videos (MP4): Transcribes audio and analyzes frames.

    Args:
        request: FastAPI request object
        file: Uploaded file

    Returns:
        HTMLResponse with the rendered PDD

    Raises:
        HTTPException: If file processing fails
    """
    temp_file_path = None
    try:
        # Create a temporary file to save the upload
        with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{file.filename}") as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name

        # Get file extension
        file_extension = Path(file.filename).suffix.lower()

        # Process based on file type
        if file_extension in ['.pdf', '.docx']:
            # Parse document to extract text
            process_text = parse_document(temp_file_path)
        elif file_extension in ['.mp4', '.mov', '.avi']:
            # Process video: transcribe audio, analyze frames, synthesize
            transcript = transcribe_audio_from_video(temp_file_path)
            visual_actions = analyze_video_frames(temp_file_path, transcript)
            process_text = synthesize_video_analysis(transcript, visual_actions)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file_extension}. "
                       "Supported types: .pdf, .docx, .mp4, .mov, .avi"
            )

        # Generate PDD from the extracted text
        sections = extract_pdd_sections(process_text)

        # Get the process name for the title - strip HTML tags
        import re
        process_name = sections[0]["content"] if sections else "Process Design Document"
        process_name_clean = re.sub('<[^<]+?>', '', process_name).strip()
        process_name = process_name_clean if process_name_clean else "Process Design Document"

        # Generate diagram from process steps
        diagram_code = None
        if len(sections) >= 6:
            process_steps = sections[5]["content"]  # Detailed Process Steps section
            try:
                diagram_code = generate_mermaid_diagram(process_steps)
            except Exception as e:
                print(f"Warning: Failed to generate diagram: {str(e)}")
                diagram_code = None

        # Setup Jinja2 environment
        template_path = Path(__file__).parent.parent / "templates"
        env = Environment(loader=FileSystemLoader(template_path))
        template = env.get_template("pdd_template.html")

        # Render the template
        html_content = template.render(
            process_name=process_name,
            sections=sections,
            diagram_code=diagram_code
        )

        return HTMLResponse(content=html_content)

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process file: {str(e)}"
        )
    finally:
        # Clean up temporary file
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except Exception:
                pass


@router.post("/refine-section")
async def refine_section(request: Request, body: RefineSectionRequest):
    """
    Refine a PDD section based on user feedback using AI.

    Args:
        request: FastAPI request object
        body: Request containing section_name, current_content, and user_feedback

    Returns:
        JSONResponse with refined content

    Raises:
        HTTPException: If refinement fails
    """
    try:
        from langchain_openai import ChatOpenAI

        # Initialize LLM
        llm = ChatOpenAI(model="gpt-4o", temperature=0)

        # Create refinement prompt
        refine_prompt = f"""
You are an expert UiPath Business Analyst refining a PDD section.

The section name is: '{body.section_name}'
The current content is: '{body.current_content}'
The user has provided the following feedback: '{body.user_feedback}'

Rewrite the section content based on the user's feedback.
- Maintain a professional tone
- Adhere to UiPath documentation standards
- Keep the content clear and concise
- Preserve the structure (bulleted lists, numbered steps, etc.) where appropriate

Output only the revised content, nothing else.
"""

        # Generate refined content
        response = llm.invoke(refine_prompt)
        refined_content = response.content.strip()

        return {"refined_content": refined_content}

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to refine section: {str(e)}"
        )


@router.post("/chat")
async def chat(request: Request, body: ChatRequest):
    """
    Clarification chat endpoint for asking questions about PDD generation.

    Args:
        request: FastAPI request object
        body: Request containing message and optional context

    Returns:
        JSONResponse with AI response

    Raises:
        HTTPException: If chat fails
    """
    try:
        from langchain_openai import ChatOpenAI

        # Initialize LLM
        llm = ChatOpenAI(model="gpt-4o", temperature=0)

        # Create chat prompt
        if body.context:
            chat_prompt = f"""
You are an expert UiPath Business Analyst helping a user create a Process Design Document (PDD).

Context about the process: {body.context}

User's question: {body.message}

Provide a helpful, concise response to assist with PDD creation.
If the question is about the process, use the provided context.
Keep responses focused on UiPath and RPA documentation standards.
"""
        else:
            chat_prompt = f"""
You are an expert UiPath Business Analyst helping a user create a Process Design Document (PDD).

User's question: {body.message}

Provide a helpful, concise response to assist with PDD creation.
Keep responses focused on UiPath and RPA documentation standards.
"""

        # Generate response
        response = llm.invoke(chat_prompt)
        ai_response = response.content.strip()

        return {"response": ai_response}

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process chat: {str(e)}"
        )


# ==================== JSON API Endpoints for Interactive UI ====================

@router.post("/api/generate-pdd-json")
async def generate_pdd_json(request: Request, body: GeneratePDDRequest):
    """
    Generate a PDD and return structured JSON (for interactive UI).

    Args:
        request: FastAPI request object
        body: Request body containing process_text

    Returns:
        JSONResponse with sections and diagram code

    Raises:
        HTTPException: If PDD generation fails
    """
    try:
        # Extract PDD sections
        sections = extract_pdd_sections(body.process_text)

        # Get the process name - strip HTML tags to get plain text
        import re
        process_name = sections[0]["content"] if sections else "Process Design Document"
        # Remove HTML tags to get clean text for title
        process_name_clean = re.sub('<[^<]+?>', '', process_name).strip()
        process_name = process_name_clean if process_name_clean else "Process Design Document"

        # Generate diagram
        diagram_code = None
        if len(sections) >= 6:
            process_steps = sections[5]["content"]
            try:
                diagram_code = generate_mermaid_diagram(process_steps)
            except Exception as e:
                print(f"Warning: Failed to generate diagram: {str(e)}")

        return {
            "process_name": process_name,
            "sections": sections,
            "diagram_code": diagram_code
        }

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate PDD: {str(e)}"
        )


@router.post("/api/upload-and-process-json")
async def upload_and_process_json(request: Request, file: UploadFile = File(...)):
    """
    Upload a file and return structured JSON (for interactive UI).

    Args:
        request: FastAPI request object
        file: Uploaded file

    Returns:
        JSONResponse with sections and diagram code

    Raises:
        HTTPException: If file processing fails
    """
    temp_file_path = None
    try:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{file.filename}") as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name

        # Get file extension
        file_extension = Path(file.filename).suffix.lower()

        # Process based on file type
        if file_extension in ['.pdf', '.docx']:
            process_text = parse_document(temp_file_path)
        elif file_extension in ['.mp4', '.mov', '.avi']:
            transcript = transcribe_audio_from_video(temp_file_path)
            visual_actions = analyze_video_frames(temp_file_path, transcript)
            process_text = synthesize_video_analysis(transcript, visual_actions)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file_extension}"
            )

        # Generate PDD
        sections = extract_pdd_sections(process_text)

        # Get the process name - strip HTML tags to get plain text
        import re
        process_name = sections[0]["content"] if sections else "Process Design Document"
        # Remove HTML tags to get clean text for title
        process_name_clean = re.sub('<[^<]+?>', '', process_name).strip()
        process_name = process_name_clean if process_name_clean else "Process Design Document"

        # Generate diagram
        diagram_code = None
        if len(sections) >= 6:
            process_steps = sections[5]["content"]
            try:
                diagram_code = generate_mermaid_diagram(process_steps)
            except Exception as e:
                print(f"Warning: Failed to generate diagram: {str(e)}")

        return {
            "process_name": process_name,
            "sections": sections,
            "diagram_code": diagram_code
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process file: {str(e)}"
        )
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except Exception:
                pass



# ==================== Export Endpoint ====================

def _create_html_document(process_name: str, sections: List[Dict], diagram_code: Optional[str]) -> str:
    """Create HTML document using the Jinja2 template."""
    template_path = Path(__file__).parent.parent / "templates"
    env = Environment(loader=FileSystemLoader(template_path))
    template = env.get_template("pdd_template.html")
    
    return template.render(
        process_name=process_name,
        sections=sections,
        diagram_code=diagram_code
    )


def _create_word_document(process_name: str, sections: List[Dict], diagram_code: Optional[str]) -> bytes:
    """Create a Word document (.docx) from the PDD data."""
    doc = Document()

    # Title
    title = doc.add_heading(process_name, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add subtitle
    subtitle = doc.add_paragraph('Process Definition Document (PDD)')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].italic = True

    # Find the section after which to insert the diagram
    diagram_insert_after = -1
    for i, section in enumerate(sections):
        if (section['name'] == "Process Overview (AS IS)" or
            section['name'] == "High Level Process Map (AS IS)" or
            section['name'] == "Detailed Process Map (AS IS)"):
            diagram_insert_after = i
            break

    # Add sections
    for i, section in enumerate(sections):
        section_name = section['name']
        content = section['content']

        # Handle different heading levels based on section name
        if 'Purpose' in section_name or 'Objectives' in section_name or 'Key Contacts' in section_name or 'Pre-requisites' in section_name:
            doc.add_heading(section_name, level=2)
        elif 'AS IS' in section_name and 'Overview' in section_name:
            doc.add_heading('AS IS Process Description', level=1)
            doc.add_heading(section_name, level=2)
        elif 'TO BE' in section_name and 'Overview' in section_name:
            doc.add_heading('TO BE Process Description', level=1)
            doc.add_heading(section_name, level=2)
        elif 'Exceptions Handling' in section_name or 'Reporting' in section_name:
            doc.add_heading(section_name, level=1)
        else:
            doc.add_heading(section_name, level=2)

        # Add content
        # Convert HTML content to plain text paragraphs
        lines = content.replace('<br>', '\n').replace('</p>', '\n').split('\n')
        current_paragraph = None

        for line in lines:
            line = line.strip()
            if not line:
                if current_paragraph:
                    current_paragraph = None
                doc.add_paragraph()
                continue

            # Remove HTML tags
            import re
            clean_line = re.sub('<[^<]+?>', '', line)

            if clean_line:
                # Check if it's a list item
                is_numbered_list = len(clean_line) >= 4 and clean_line[:3].isdigit() and clean_line[3] == '.'
                if clean_line.startswith(('-', '*', '•')) or is_numbered_list:
                    doc.add_paragraph(clean_line, style='List Bullet')
                else:
                    doc.add_paragraph(clean_line)

        # Insert diagram after the specified section
        if diagram_code and i == diagram_insert_after:
            doc.add_heading('Process Flow Diagram', level=2)
            doc.add_paragraph('Diagram code (Mermaid syntax):')
            doc.add_paragraph(diagram_code)
            doc.add_paragraph()  # Empty line

    # Save to bytes
    from io import BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    return doc_io.getvalue()


@router.post("/api/export-pdd")
async def export_pdd(request: Request, body: ExportPDDRequest):
    """
    Export PDD as PDF, Word, or HTML document.

    Args:
        request: FastAPI request object
        body: Request containing PDD data and format

    Returns:
        StreamingResponse with the document file

    Raises:
        HTTPException: If export fails
    """
    try:
        if body.format == 'html':
            # Generate HTML
            html_content = _create_html_document(body.process_name, body.sections, body.diagram_code)
            
            # Return as downloadable file
            return StreamingResponse(
                iter([html_content]),
                media_type="text/html",
                headers={
                    "Content-Disposition": f"attachment; filename=PDD_{body.process_name.replace(' ', '_')}.html"
                }
            )
        
        elif body.format == 'docx':
            # Generate Word document
            doc_bytes = _create_word_document(body.process_name, body.sections, body.diagram_code)
            
            return StreamingResponse(
                iter([doc_bytes]),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename=PDD_{body.process_name.replace(' ', '_')}.docx"
                }
            )
        
        elif body.format == 'pdf':
            # For PDF, we return HTML that can be printed to PDF
            # This is a workaround since true PDF generation requires additional libraries
            html_content = _create_html_document(body.process_name, body.sections, body.diagram_code)

            return HTMLResponse(
                content=html_content,
                headers={
                    "Content-Disposition": f"inline; filename=PDD_{body.process_name.replace(' ', '_')}.html"
                }
            )
        
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported format: {body.format}. Supported: html, docx, pdf"
            )
    
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to export document: {str(e)}"
        )
