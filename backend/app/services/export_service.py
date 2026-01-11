"""
Export Service for generating documents in various formats.

This service handles the export of PDD data to different document formats
including HTML, Word (DOCX), and PDF. It encapsulates the document generation
logic, making it easier to maintain and test.
"""

from io import BytesIO
from pathlib import Path
from typing import Optional, List, Dict

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from jinja2 import Environment, FileSystemLoader


class ExportService:
    """
    Service for exporting PDD documents.

    This service provides methods to export PDD data to various formats:
    - HTML: Complete HTML document with embedded styles
    - DOCX: Microsoft Word document
    - PDF: HTML wrapper (PDF generation requires additional libraries)

    Attributes:
        template_path: Path to Jinja2 templates directory
        _env: Lazy-loaded Jinja2 Environment

    Example:
        >>> service = ExportService()
        >>> # Export to HTML
        >>> html = service.create_html_document(name, sections, diagram)
        >>> # Export to Word
        >>> doc_bytes = service.create_word_document(name, sections, diagram)
    """

    def __init__(self, template_path: Optional[Path] = None):
        """
        Initialize ExportService.

        Args:
            template_path: Path to templates directory. Defaults to app/templates
        """
        if template_path is None:
            template_path = Path(__file__).parent.parent / "templates"
        self.template_path = template_path
        self._env = None

    @property
    def env(self) -> Environment:
        """
        Lazy-loaded Jinja2 Environment.

        Returns:
            Configured Jinja2 Environment

        Example:
            >>> service = ExportService()
            >>> env = service.env
            >>> template = env.get_template("pdd_template.html")
        """
        if self._env is None:
            self._env = Environment(loader=FileSystemLoader(self.template_path))
        return self._env

    def create_html_document(
        self,
        process_name: str,
        sections: List[Dict],
        diagram_code: Optional[str]
    ) -> str:
        """
        Create HTML document using the Jinja2 template.

        Renders PDD data into a complete HTML document with embedded
        styles and Mermaid diagram support.

        Args:
            process_name: Clean process name for title
            sections: List of section dictionaries
            diagram_code: Mermaid diagram code or None

        Returns:
            Complete HTML document as string

        Example:
            >>> service = ExportService()
            >>> html = service.create_html_document(
            ...     "Invoice Process",
            ...     sections,
            ...     diagram_code
            ... )
            >>> # Can save to file or return as HTTP response
        """
        template = self.env.get_template("pdd_template.html")

        return template.render(
            process_name=process_name,
            sections=sections,
            diagram_code=diagram_code
        )

    def create_word_document(
        self,
        process_name: str,
        sections: List[Dict],
        diagram_code: Optional[str]
    ) -> bytes:
        """
        Create a Word document (.docx) from PDD data.

        Generates a properly formatted Microsoft Word document with
        headings, paragraphs, lists, and diagram code.

        Args:
            process_name: Clean process name for title
            sections: List of section dictionaries
            diagram_code: Mermaid diagram code or None

        Returns:
            Word document as bytes

        Example:
            >>> service = ExportService()
            >>> doc_bytes = service.create_word_document(
            ...     "Invoice Process",
            ...     sections,
            ...     diagram_code
            ... )
            >>> # Can save as .docx file or return as HTTP response
        """
        doc = Document()

        # Title
        title = doc.add_heading(process_name, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add subtitle
        subtitle = doc.add_paragraph('Process Definition Document (PDD)')
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle.runs[0].italic = True

        # Find the section after which to insert the diagram
        diagram_insert_after = -1
        for i, section in enumerate(sections):
            if (section['name'] == "Process Overview (AS IS)" or
                section['name'] == "High Level Process Map (AS IS)" or
                section['name'] == "Detailed Process Map (AS IS)"):
                diagram_insert_after = i
                break

        # Add sections
        for i, section in enumerate(sections):
            section_name = section['name']
            content = section['content']

            # Handle different heading levels based on section name
            if ('Purpose' in section_name or 'Objectives' in section_name or
                'Key Contacts' in section_name or 'Pre-requisites' in section_name):
                doc.add_heading(section_name, level=2)
            elif 'AS IS' in section_name and 'Overview' in section_name:
                doc.add_heading('AS IS Process Description', level=1)
                doc.add_heading(section_name, level=2)
            elif 'TO BE' in section_name and 'Overview' in section_name:
                doc.add_heading('TO BE Process Description', level=1)
                doc.add_heading(section_name, level=2)
            elif 'Exceptions Handling' in section_name or 'Reporting' in section_name:
                doc.add_heading(section_name, level=1)
            else:
                doc.add_heading(section_name, level=2)

            # Add content
            # Convert HTML content to plain text paragraphs
            lines = content.replace('<br>', '\n').replace('</p>', '\n').split('\n')
            current_paragraph = None

            for line in lines:
                line = line.strip()
                if not line:
                    if current_paragraph:
                        current_paragraph = None
                    doc.add_paragraph()
                    continue

                # Remove HTML tags
                import re
                clean_line = re.sub('<[^<]+?>', '', line)

                if clean_line:
                    # Check if it's a list item
                    is_numbered_list = (len(clean_line) >= 4 and
                                       clean_line[:3].isdigit() and
                                       clean_line[3] == '.')
                    if clean_line.startswith(('-', '*', '•')) or is_numbered_list:
                        doc.add_paragraph(clean_line, style='List Bullet')
                    else:
                        doc.add_paragraph(clean_line)

            # Insert diagram after the specified section
            if diagram_code and i == diagram_insert_after:
                doc.add_heading('Process Flow Diagram', level=2)
                doc.add_paragraph('Diagram code (Mermaid syntax):')
                doc.add_paragraph(diagram_code)
                doc.add_paragraph()  # Empty line

        # Save to bytes
        doc_io = BytesIO()
        doc.save(doc_io)
        return doc_io.getvalue()
